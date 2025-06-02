#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Python Word操作性能基准测试

使用python-docx库进行Word文档操作的性能测试，
与Golang WordZero库进行对比。
"""

import os
import sys
import time
import json
import tracemalloc
import psutil
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

import pytest
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 确保输出目录存在
OUTPUT_DIR = Path("../results/python")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# 统一的测试配置，与其他语言保持一致
TEST_ITERATIONS = {
    "basic": 50,       # 基础文档创建
    "complex": 30,     # 复杂格式化
    "table": 20,       # 表格操作
    "largeTable": 10,  # 大表格处理
    "largeDoc": 5,     # 大型文档
    "memory": 10,      # 内存使用测试
}


class PerformanceTester:
    """性能测试工具类"""
    
    def __init__(self):
        self.results = []
        
    def run_test(self, name: str, test_function, iterations: int = 10):
        """运行性能测试"""
        print(f"\n开始测试: {name}")
        times = []
        
        for i in range(iterations):
            start_time = time.perf_counter()
            test_function(i)
            end_time = time.perf_counter()
            
            duration = (end_time - start_time) * 1000  # 转换为毫秒
            times.append(duration)
            
            if i % max(1, iterations // 10) == 0:
                print(f"  进度: {i + 1}/{iterations}")
        
        avg_time = sum(times) / len(times)
        min_time = min(times)
        max_time = max(times)
        
        result = {
            'name': name,
            'avgTime': round(avg_time, 2),
            'minTime': round(min_time, 2),
            'maxTime': round(max_time, 2),
            'iterations': iterations
        }
        
        self.results.append(result)
        print(f"  平均耗时: {result['avgTime']}ms")
        print(f"  最小耗时: {result['minTime']}ms")
        print(f"  最大耗时: {result['maxTime']}ms")
        
        return result
    
    def generate_report(self):
        """生成性能测试报告"""
        print('\n=== Python 性能测试报告 ===')
        for result in self.results:
            print(f"{result['name']}: {result['avgTime']}ms (平均)")
        
        # 保存详细报告
        report_path = OUTPUT_DIR / 'performance_report.json'
        report_data = {
            'timestamp': datetime.now().isoformat(),
            'platform': 'Python',
            'pythonVersion': f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}",
            'results': self.results
        }
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False)
        
        print(f"\n详细报告已保存到: {report_path}")


def test_basic_document_creation(index: int):
    """基础文档创建测试"""
    doc = Document()
    doc.add_paragraph("这是一个基础性能测试文档")
    doc.add_paragraph("测试内容包括基本的文本添加功能")
    
    filename = OUTPUT_DIR / f"basic_doc_{index}.docx"
    doc.save(filename)


def test_complex_formatting(index: int):
    """复杂格式化测试"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading("性能测试报告", level=1)
    subtitle = doc.add_heading("测试概述", level=2)
    
    # 添加格式化文本
    para = doc.add_paragraph()
    
    # 粗体文本
    run = para.add_run("粗体文本")
    run.bold = True
    
    para.add_run(" ")
    
    # 斜体文本
    run = para.add_run("斜体文本")
    run.italic = True
    
    para.add_run(" ")
    
    # 彩色文本
    run = para.add_run("彩色文本")
    run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
    
    # 添加多个格式化段落
    for j in range(10):
        para = doc.add_paragraph(f"这是第{j + 1}个段落，包含复杂格式化")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    filename = OUTPUT_DIR / f"complex_formatting_{index}.docx"
    doc.save(filename)


def test_table_operations(index: int):
    """表格操作测试"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("表格性能测试", level=1)
    
    # 创建10行5列的表格
    table = doc.add_table(rows=10, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 填充表格数据
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f"R{row_idx + 1}C{col_idx + 1}"
    
    filename = OUTPUT_DIR / f"table_operations_{index}.docx"
    doc.save(filename)


def test_large_table(index: int):
    """大表格测试"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("大表格性能测试", level=1)
    
    # 创建100行10列的大表格
    table = doc.add_table(rows=100, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 填充表格数据
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f"数据_{row_idx + 1}_{col_idx + 1}"
    
    filename = OUTPUT_DIR / f"large_table_{index}.docx"
    doc.save(filename)


def test_large_document(index: int):
    """大型文档测试"""
    doc = Document()
    
    # 添加主标题
    doc.add_heading("大型文档性能测试", level=1)
    
    # 添加1000个段落
    for j in range(1000):
        if j % 10 == 0:
            # 每10个段落添加一个标题
            doc.add_heading(f"章节 {j // 10 + 1}", level=2)
        
        doc.add_paragraph(
            f"这是第{j + 1}个段落。Lorem ipsum dolor sit amet, consectetur "
            f"adipiscing elit. Sed do eiusmod tempor incididunt ut labore et "
            f"dolore magna aliqua."
        )
    
    # 添加一个中等大小的表格
    table = doc.add_table(rows=20, cols=8)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f"表格数据{row_idx + 1}-{col_idx + 1}"
    
    filename = OUTPUT_DIR / f"large_document_{index}.docx"
    doc.save(filename)


def test_memory_usage(index: int):
    """内存使用测试"""
    # 开始内存追踪
    tracemalloc.start()
    initial_memory = tracemalloc.get_traced_memory()[0]
    
    doc = Document()
    
    # 创建复杂内容
    for j in range(100):
        doc.add_paragraph(f"段落{j + 1}: 这是一个测试段落，用于测试内存使用情况")
    
    # 添加表格
    table = doc.add_table(rows=50, cols=6)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = f"单元格{row_idx + 1}-{col_idx + 1}"
    
    # 保存文档
    filename = OUTPUT_DIR / f"memory_test_{index}.docx"
    doc.save(filename)
    
    # 检查内存使用
    final_memory = tracemalloc.get_traced_memory()[0]
    if index == 0:
        memory_used = (final_memory - initial_memory) / 1024  # 转换为KB
        print(f"  内存使用: {memory_used:.0f}KB")
    
    tracemalloc.stop()


# pytest benchmark 测试函数
def test_benchmark_basic_document(benchmark):
    """pytest-benchmark: 基础文档创建"""
    benchmark(test_basic_document_creation, 0)


def test_benchmark_complex_formatting(benchmark):
    """pytest-benchmark: 复杂格式化"""
    benchmark(test_complex_formatting, 0)


def test_benchmark_table_operations(benchmark):
    """pytest-benchmark: 表格操作"""
    benchmark(test_table_operations, 0)


def test_benchmark_large_table(benchmark):
    """pytest-benchmark: 大表格"""
    benchmark(test_large_table, 0)


def test_benchmark_large_document(benchmark):
    """pytest-benchmark: 大型文档"""
    benchmark(test_large_document, 0)


def test_benchmark_memory_usage(benchmark):
    """pytest-benchmark: 内存使用"""
    benchmark(test_memory_usage, 0)


class TestPerformanceComparison:
    """性能对比测试类"""
    
    def test_performance_comparison(self):
        """运行完整的性能对比测试"""
        print('开始 Python (python-docx库) 性能基准测试...')
        print(f'Python 版本: {sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}')
        print(f'输出目录: {OUTPUT_DIR}')
        
        tester = PerformanceTester()
        
        # 运行各项测试（使用统一的迭代次数配置）
        tester.run_test('基础文档创建', test_basic_document_creation, TEST_ITERATIONS["basic"])
        tester.run_test('复杂格式化', test_complex_formatting, TEST_ITERATIONS["complex"])
        tester.run_test('表格操作', test_table_operations, TEST_ITERATIONS["table"])
        tester.run_test('大表格处理', test_large_table, TEST_ITERATIONS["largeTable"])
        tester.run_test('大型文档', test_large_document, TEST_ITERATIONS["largeDoc"])
        tester.run_test('内存使用测试', test_memory_usage, TEST_ITERATIONS["memory"])
        
        # 生成报告
        tester.generate_report()


def main():
    """主函数，直接运行时执行"""
    print("开始 Python 性能基准测试...")
    
    tester = PerformanceTester()
    
    try:
        # 使用统一的迭代次数配置
        tester.run_test('基础文档创建', test_basic_document_creation, TEST_ITERATIONS["basic"])
        tester.run_test('复杂格式化', test_complex_formatting, TEST_ITERATIONS["complex"])
        tester.run_test('表格操作', test_table_operations, TEST_ITERATIONS["table"])
        tester.run_test('大表格处理', test_large_table, TEST_ITERATIONS["largeTable"])
        tester.run_test('大型文档', test_large_document, TEST_ITERATIONS["largeDoc"])
        tester.run_test('内存使用测试', test_memory_usage, TEST_ITERATIONS["memory"])
        
        tester.generate_report()
        
        print('\n所有测试完成！')
        
    except Exception as e:
        print(f'测试过程中发生错误: {e}')
        raise


if __name__ == '__main__':
    main() 